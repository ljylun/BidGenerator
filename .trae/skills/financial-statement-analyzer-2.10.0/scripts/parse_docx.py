#!/usr/bin/env python3
"""Word文档解析模块 — 使用 python-docx 解析 .docx 中的财务报表表格。

提取所有表格并通过关键词分类到三表（BS/IS/CF），处理合并单元格。

作者: 优方皑尔 Uform Ai
版本: v1.1.0
"""

from __future__ import annotations

import logging
import re
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# 三表关键词 — 与项目其他模块保持一致
# ---------------------------------------------------------------------------

_BALANCE_SHEET_KEYWORDS: List[str] = [
    "资产负债表", "资产总计", "负债总计", "所有者权益",
    "流动资产", "非流动资产", "流动负债", "非流动负债",
    "balance sheet", "total assets", "total liabilities",
    "equity",
]

_INCOME_STATEMENT_KEYWORDS: List[str] = [
    "利润表", "营业收入", "营业成本", "净利润", "利润总额",
    "营业利润", "销售费用", "管理费用", "综合收益",
    "income statement", "revenue", "net income",
    "gross profit",
]

_CASH_FLOW_KEYWORDS: List[str] = [
    "现金流量表", "经营活动产生", "投资活动产生", "筹资活动产生",
    "现金及现金等价物", "期初现金", "期末现金",
    "cash flow", "operating activities",
    "investing activities", "financing activities",
]


def _resolve_merged_cells(table: Any) -> List[List[str]]:
    """处理 Word 表格中的合并单元格。

    对于已合并的单元格，将其内容复制到所有被合并的虚拟位置上。

    Args:
        table: python-docx 的 Table 对象。

    Returns:
        List[List[str]]: 完整的、展开合并单元格后的二维表格数据。
    """
    import copy

    rows_data: List[List[str]] = []
    n_rows: int = len(table.rows)
    n_cols: int = len(table.columns)

    # 初始化空矩阵
    cell_matrix: List[List[Optional[str]]] = [
        [None for _ in range(n_cols)] for _ in range(n_rows)
    ]

    # 标记已被合并单元格占用的位置
    for row_idx, row in enumerate(table.rows):
        col_idx: int = 0
        for cell in row.cells:
            # 跳过已被合并单元格占用的位置
            while col_idx < n_cols and cell_matrix[row_idx][col_idx] is not None:
                col_idx += 1
            if col_idx >= n_cols:
                break

            # 获取合并信息
            # python-docx 中通过 _tc 元素获取 gridSpan / vMerge
            tc = cell._tc
            tcPr = tc.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr"
            )

            grid_span: int = 1
            v_merge: Optional[str] = None  # 'restart' | 'continue' | None

            if tcPr is not None:
                # 水平合并
                grid_span_el = tcPr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan"
                )
                if grid_span_el is not None:
                    grid_span = int(grid_span_el.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "1"
                    ))

                # 垂直合并
                v_merge_el = tcPr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge"
                )
                if v_merge_el is not None:
                    v_merge_val = v_merge_el.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                    )
                    v_merge = "restart" if v_merge_val == "restart" else "continue"

            # 获取单元格文本
            cell_text: str = cell.text.strip()

            # 处理垂直合并
            if v_merge == "restart":
                # 记录合并起始位置和文本
                v_merge_text: str = cell_text
            elif v_merge == "continue":
                # 延续之前的合并，使用之前记录的文本
                cell_text = v_merge_text if "v_merge_text" in dir() else cell_text

            # 填充矩阵
            for span_col in range(grid_span):
                target_col: int = col_idx + span_col
                if target_col < n_cols and row_idx < n_rows:
                    cell_matrix[row_idx][target_col] = cell_text

            col_idx += grid_span

    # 转换为字符串列表
    for row in cell_matrix:
        rows_data.append([cell or "" for cell in row])

    return rows_data


def _table_to_dataframe(
    table_data: List[List[str]],
) -> Optional[pd.DataFrame]:
    """将二维表格数据转换为 pandas DataFrame。

    自动识别表头行：
        - 若第一行含数值则生成默认列名
        - 否则以第一行为表头

    Args:
        table_data: 二维字符串列表。

    Returns:
        pd.DataFrame 或 None（数据不足时）。
    """
    if not table_data or len(table_data) < 2:
        return None

    # 清理全空行
    cleaned: List[List[str]] = [
        row for row in table_data
        if any(cell.strip() for cell in row)
    ]
    if len(cleaned) < 2:
        return None

    # 标准化列数
    max_cols: int = max(len(r) for r in cleaned)
    for row in cleaned:
        while len(row) < max_cols:
            row.append("")

    # 判断第一行是否为表头
    first_row: List[str] = cleaned[0]
    numeric_count: int = sum(
        1 for cell in first_row if _is_numeric(cell)
    )
    if numeric_count >= len(first_row) // 2:
        # 第一行是数据
        columns: List[str] = ["科目"] + [
            f"期间{i+1}" for i in range(max_cols - 1)
        ]
        data_rows: List[List[str]] = cleaned
    else:
        columns = [str(c).strip() for c in first_row]
        data_rows = cleaned[1:]

    # 确保列名唯一
    seen: Dict[str, int] = {}
    unique_cols: List[str] = []
    for c in columns:
        name = c or "col"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        unique_cols.append(name)

    try:
        df: pd.DataFrame = pd.DataFrame(data_rows, columns=unique_cols)
    except Exception:
        return None

    df = df.dropna(how="all")
    return df


def _classify_table(df: pd.DataFrame) -> str:
    """通过关键词分类 DataFrame 属于哪个报表。

    Args:
        df: 数据 DataFrame。

    Returns:
        str: 'balance_sheet' | 'income_statement' | 'cash_flow' | 'unknown'
    """
    # 在全表中搜索关键词
    all_text: str = " ".join(
        str(v).strip() for col in df.columns for v in df[col].dropna().head(30)
    )

    bs_score: int = sum(
        1 for kw in _BALANCE_SHEET_KEYWORDS if kw.lower() in all_text.lower()
    )
    is_score: int = sum(
        1 for kw in _INCOME_STATEMENT_KEYWORDS if kw.lower() in all_text.lower()
    )
    cf_score: int = sum(
        1 for kw in _CASH_FLOW_KEYWORDS if kw.lower() in all_text.lower()
    )

    scores: List[Tuple[str, int]] = [
        ("balance_sheet", bs_score),
        ("income_statement", is_score),
        ("cash_flow", cf_score),
    ]
    scores.sort(key=lambda x: x[1], reverse=True)

    if scores[0][1] > 0:
        return scores[0][0]
    return "unknown"


def _is_numeric(text: str) -> bool:
    """判断文本是否为数值。

    Args:
        text: 输入文本。

    Returns:
        bool: 是否为数值。
    """
    text = text.strip().replace(",", "").replace(" ", "").replace("%", "")
    text = text.replace("(", "").replace(")", "")
    text = re.sub(r"[¥€$万亿千百十元整]", "", text)
    if not text:
        return False
    try:
        float(text)
        return True
    except ValueError:
        return False


def _extract_docx_text(doc: Any) -> str:
    """提取 DOCX 文档的全部文本内容。

    Args:
        doc: python-docx 的 Document 对象。

    Returns:
        str: 文档全文。
    """
    paragraphs: List[str] = []
    for para in doc.paragraphs:
        text: str = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------


def parse_docx(filepath: str) -> Dict[str, Any]:
    """解析 Word 文档中的财务报表。

    提取所有表格，通过关键词自动分类到三表，处理合并单元格。

    Args:
        filepath: .docx 文件路径。

    Returns:
        dict: 统一输出格式:
            {
                "balance_sheet": pd.DataFrame | None,
                "income_statement": pd.DataFrame | None,
                "cash_flow": pd.DataFrame | None,
                "metadata": {
                    "source": str,
                    "source_format": "docx",
                    "extraction_method": "direct",
                    "extraction_confidence": float,
                    "warnings": [...],
                    "audit_report_detected": bool,
                    ...
                }
            }

        解析失败时返回部分数据 + warnings，不抛异常。
    """
    warnings: List[str] = []
    metadata: Dict[str, Any] = {
        "source": filepath,
        "source_format": "docx",
        "extraction_method": "direct",
        "extraction_confidence": 0.0,
        "warnings": [],
        "audit_report_detected": False,
        "total_tables": 0,
    }

    bs_df: Optional[pd.DataFrame] = None
    is_df: Optional[pd.DataFrame] = None
    cf_df: Optional[pd.DataFrame] = None

    try:
        from docx import Document
    except ImportError:
        warnings.append("python-docx 未安装。请执行: pip install python-docx")
        metadata["warnings"] = warnings
        return {
            "balance_sheet": None,
            "income_statement": None,
            "cash_flow": None,
            "metadata": metadata,
        }

    # 打开文档
    try:
        doc = Document(filepath)
    except Exception as exc:
        warnings.append(f"无法打开 Word 文档 '{filepath}': {exc}")
        metadata["warnings"] = warnings
        return {
            "balance_sheet": None,
            "income_statement": None,
            "cash_flow": None,
            "metadata": metadata,
        }

    tables = doc.tables
    metadata["total_tables"] = len(tables)

    if not tables:
        warnings.append("Word 文档中未找到任何表格。")
        metadata["warnings"] = warnings
        return {
            "balance_sheet": None,
            "income_statement": None,
            "cash_flow": None,
            "metadata": metadata,
        }

    # 提取全文用于审计报告检测
    full_text: str = _extract_docx_text(doc)

    # 审计报告检测
    try:
        from .parse_audit_report import is_audit_report as _is_audit
    except ImportError:
        # 如果无法导入，使用内联检测
        _is_audit = None

    if _is_audit is not None:
        try:
            if _is_audit(full_text):
                metadata["audit_report_detected"] = True
                logger.info("Detected audit report in DOCX, delegating to parse_audit_report.")
                try:
                    from .parse_audit_report import parse_audit_report_from_docx
                    result = parse_audit_report_from_docx(filepath)
                    # 合并 metadata
                    merged_meta = {**metadata, **result.get("metadata", {})}
                    result["metadata"] = merged_meta
                    return result
                except ImportError:
                    warnings.append(
                        "检测到审计报告但无法加载审计报告解析器，"
                        "使用标准 DOCX 解析。"
                    )
        except Exception as exc:
            logger.warning("Audit report check failed: %s", exc)

    # 处理每个表格
    classified: Dict[str, pd.DataFrame] = {}
    for idx, table in enumerate(tables):
        try:
            table_data: List[List[str]] = _resolve_merged_cells(table)
        except Exception as exc:
            logger.warning("Failed to resolve merged cells in table %d: %s", idx, exc)
            continue

        df: Optional[pd.DataFrame] = _table_to_dataframe(table_data)
        if df is None or df.empty:
            continue

        table_type: str = _classify_table(df)

        if table_type == "balance_sheet" and "balance_sheet" not in classified:
            classified["balance_sheet"] = df
        elif table_type == "income_statement" and "income_statement" not in classified:
            classified["income_statement"] = df
        elif table_type == "cash_flow" and "cash_flow" not in classified:
            classified["cash_flow"] = df

    bs_df = classified.get("balance_sheet")
    is_df = classified.get("income_statement")
    cf_df = classified.get("cash_flow")

    # 计算置信度
    found_tables: int = sum(
        1 for df in [bs_df, is_df, cf_df] if df is not None
    )
    confidence: float = (found_tables / 3.0) * 100.0
    if found_tables < 3 and len(tables) > 0:
        # 有表格但没找到三表，可能是非标准格式
        confidence = min(confidence, 60.0)
    metadata["extraction_confidence"] = round(confidence, 1)

    if found_tables == 0:
        warnings.append(
            "未能从 Word 表格中识别出资产负债表、利润表或现金流量表。"
            "请确认文档包含标准财务报表表格。"
        )

    metadata["warnings"] = warnings
    metadata["found_tables"] = found_tables

    return {
        "balance_sheet": bs_df,
        "income_statement": is_df,
        "cash_flow": cf_df,
        "metadata": metadata,
    }
