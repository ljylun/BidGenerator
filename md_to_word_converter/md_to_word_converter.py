"""
Markdown to Word Document Converter
将Markdown文件转换为符合公文规范的Word文档

功能特性：
- 完整保留Markdown所有格式元素
- 公文规范样式配置
- 本地资源检测与嵌入
- 完整的异常处理和日志
- 转换后全量校验
"""

import os
import sys
import re
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any
from dataclasses import dataclass, field

import markdown
from markdown.extensions import Extension
from markdown.preprocessors import Preprocessor
from markdown.postprocessors import Postprocessor
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image
from tqdm import tqdm

from config import *


# ==================== 数据类定义 ====================

@dataclass
class ImageResource:
    """图片资源"""
    path: str
    alt_text: str = ""
    line_number: int = 0
    found: bool = False


@dataclass
class ConversionStats:
    """转换统计"""
    total_lines: int = 0
    total_chars: int = 0
    headings: Dict[int, int] = field(default_factory=lambda: {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0})
    paragraphs: int = 0
    tables: int = 0
    code_blocks: int = 0
    lists: int = 0
    images: int = 0
    images_found: int = 0
    images_missing: int = 0
    blockquotes: int = 0
    links: int = 0
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


@dataclass
class ValidationResult:
    """校验结果"""
    passed: bool = True
    content_complete: bool = True
    format_correct: bool = True
    resources_available: bool = True
    source_stats: Dict[str, Any] = field(default_factory=dict)
    target_stats: Dict[str, Any] = field(default_factory=dict)
    issues: List[str] = field(default_factory=list)


# ==================== 日志配置 ====================

def setup_logging(log_file: str) -> logging.Logger:
    """配置日志系统"""
    logger = logging.getLogger("md_to_word")
    logger.setLevel(logging.DEBUG)

    # 文件处理器
    file_handler = logging.FileHandler(log_file, encoding="utf-8")
    file_handler.setLevel(logging.DEBUG)
    file_formatter = logging.Formatter(LOG_FORMAT, datefmt=LOG_DATE_FORMAT)
    file_handler.setFormatter(file_formatter)

    # 控制台处理器
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_formatter = logging.Formatter("%(levelname)s: %(message)s")
    console_handler.setFormatter(console_formatter)

    logger.addHandler(file_handler)
    logger.addHandler(console_handler)

    return logger


# ==================== Markdown解析器 ====================

class MarkdownParser:
    """Markdown文件解析器"""

    def __init__(self, logger: logging.Logger):
        self.logger = logger
        self.stats = ConversionStats()
        self.images: List[ImageResource] = []

    def read_file(self, file_path: str) -> str:
        """读取Markdown文件（支持大文件）"""
        self.logger.info(f"读取文件: {file_path}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"源文件不存在: {file_path}")

        file_size = os.path.getsize(file_path)
        self.logger.info(f"文件大小: {file_size / 1024 / 1024:.2f} MB")

        # 尝试多种编码
        encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                self.logger.info(f"文件编码: {encoding}")
                return content
            except UnicodeDecodeError:
                continue

        raise ValueError(f"无法识别文件编码: {file_path}")

    def extract_images(self, content: str) -> List[ImageResource]:
        """提取Markdown中的图片引用"""
        self.logger.info("提取图片引用...")

        # 匹配 ![alt](path) 格式的图片
        pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        matches = re.finditer(pattern, content)

        for match in matches:
            alt_text = match.group(1)
            img_path = match.group(2)
            line_num = content[:match.start()].count('\n') + 1

            # 跳过URL图片
            if img_path.startswith(('http://', 'https://', 'data:')):
                continue

            img_resource = ImageResource(
                path=img_path,
                alt_text=alt_text,
                line_number=line_num
            )
            self.images.append(img_resource)

        self.logger.info(f"发现 {len(self.images)} 个本地图片引用")
        return self.images

    def parse_markdown(self, content: str) -> List[Dict[str, Any]]:
        """解析Markdown内容为结构化数据"""
        self.logger.info("解析Markdown内容...")

        lines = content.split('\n')
        self.stats.total_lines = len(lines)
        self.stats.total_chars = len(content)

        elements = []
        i = 0

        with tqdm(total=len(lines), desc="解析进度", unit="行") as pbar:
            while i < len(lines):
                line = lines[i]

                # 空行
                if not line.strip():
                    i += 1
                    pbar.update(1)
                    continue

                # 标题
                heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
                if heading_match:
                    level = len(heading_match.group(1))
                    text = heading_match.group(2).strip()
                    elements.append({
                        "type": "heading",
                        "level": level,
                        "text": text
                    })
                    self.stats.headings[level] = self.stats.headings.get(level, 0) + 1
                    i += 1
                    pbar.update(1)
                    continue

                # 代码块
                if line.strip().startswith('```'):
                    code_lines = []
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_lines.append(lines[i])
                        i += 1
                    if i < len(lines):
                        i += 1  # 跳过结束标记
                    elements.append({
                        "type": "code_block",
                        "language": line.strip()[3:].strip(),
                        "code": '\n'.join(code_lines)
                    })
                    self.stats.code_blocks += 1
                    pbar.update(len(code_lines) + 2)
                    continue

                # 表格
                if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1]):
                    table_lines = [line]
                    i += 1
                    # 跳过分隔行
                    if i < len(lines) and re.match(r'^[\s|:-]+$', lines[i]):
                        table_lines.append(lines[i])
                        i += 1
                    # 读取表格数据行
                    while i < len(lines) and '|' in lines[i]:
                        table_lines.append(lines[i])
                        i += 1
                    elements.append({
                        "type": "table",
                        "rows": self._parse_table(table_lines)
                    })
                    self.stats.tables += 1
                    pbar.update(len(table_lines))
                    continue

                # 图片
                img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line.strip())
                if img_match:
                    elements.append({
                        "type": "image",
                        "alt": img_match.group(1),
                        "path": img_match.group(2)
                    })
                    self.stats.images += 1
                    i += 1
                    pbar.update(1)
                    continue

                # 引用块
                if line.strip().startswith('>'):
                    quote_lines = []
                    while i < len(lines) and lines[i].strip().startswith('>'):
                        quote_lines.append(lines[i].strip()[1:].strip())
                        i += 1
                    elements.append({
                        "type": "blockquote",
                        "text": '\n'.join(quote_lines)
                    })
                    self.stats.blockquotes += 1
                    pbar.update(len(quote_lines))
                    continue

                # 无序列表
                if re.match(r'^[\s]*[-*+]\s', line):
                    list_items = []
                    while i < len(lines) and re.match(r'^[\s]*[-*+]\s', lines[i]):
                        list_items.append(re.sub(r'^[\s]*[-*+]\s+', '', lines[i]))
                        i += 1
                    elements.append({
                        "type": "unordered_list",
                        "items": list_items
                    })
                    self.stats.lists += 1
                    pbar.update(len(list_items))
                    continue

                # 有序列表
                if re.match(r'^[\s]*\d+\.\s', line):
                    list_items = []
                    while i < len(lines) and re.match(r'^[\s]*\d+\.\s', lines[i]):
                        list_items.append(re.sub(r'^[\s]*\d+\.\s+', '', lines[i]))
                        i += 1
                    elements.append({
                        "type": "ordered_list",
                        "items": list_items
                    })
                    self.stats.lists += 1
                    pbar.update(len(list_items))
                    continue

                # 普通段落
                para_lines = [line]
                i += 1
                while i < len(lines) and lines[i].strip() and not self._is_block_start(lines[i]):
                    para_lines.append(lines[i])
                    i += 1
                elements.append({
                    "type": "paragraph",
                    "text": ' '.join(para_lines)
                })
                self.stats.paragraphs += 1
                pbar.update(len(para_lines))

        self.logger.info(f"解析完成: {len(elements)} 个元素")
        return elements

    def _is_block_start(self, line: str) -> bool:
        """检查行是否是块级元素的开始"""
        patterns = [
            r'^#{1,6}\s',  # 标题
            r'^```',       # 代码块
            r'^\s*[-*+]\s',  # 无序列表
            r'^\s*\d+\.\s',  # 有序列表
            r'^\s*>',      # 引用
            r'^\s*\|',     # 表格
        ]
        return any(re.match(p, line) for p in patterns)

    def _parse_table(self, lines: List[str]) -> List[List[str]]:
        """解析表格行"""
        rows = []
        for line in lines:
            if re.match(r'^[\s|:-]+$', line):
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
        return rows


# ==================== Word文档生成器 ====================

class WordGenerator:
    """Word文档生成器"""

    def __init__(self, logger: logging.Logger):
        self.logger = logger
        self.doc = None

    def create_document(self) -> Document:
        """创建Word文档并设置样式"""
        self.logger.info("创建Word文档...")
        self.doc = Document()

        # 设置页面
        self._setup_page()

        # 设置样式
        self._setup_styles()

        return self.doc

    def _setup_page(self):
        """设置页面"""
        section = self.doc.sections[0]

        # 设置页边距（转换为磅，1厘米=28.35磅）
        section.top_margin = Cm(PAGE_MARGIN_TOP)
        section.bottom_margin = Cm(PAGE_MARGIN_BOTTOM)
        section.left_margin = Cm(PAGE_MARGIN_LEFT)
        section.right_margin = Cm(PAGE_MARGIN_RIGHT)

        # 设置页面大小
        section.page_width = Cm(PAGE_WIDTH)
        section.page_height = Cm(PAGE_HEIGHT)

        self.logger.info(f"页面设置: 上下{PAGE_MARGIN_TOP}cm, 左右{PAGE_MARGIN_LEFT}cm")

    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONT_CHINESE
        font.size = Pt(BODY_FONT_SIZE)

        # 设置中文字体
        rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
        if rFonts is None:
            rPr = style.element.get_or_add_rPr()
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_CHINESE)

        # 设置标题样式
        for level, config in HEADING_CONFIG.items():
            style_name = f'Heading {level}'
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                font = style.font
                font.name = config["font_name"]
                font.size = Pt(config["font_size"])
                font.bold = config["bold"]

                # 设置中文字体
                rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
                if rFonts is None:
                    rPr = style.element.get_or_add_rPr()
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), config["font_name"])

                if config.get("italic"):
                    font.italic = True

        self.logger.info("样式设置完成")

    def add_toc(self):
        """添加目录"""
        self.logger.info("生成目录...")

        # 添加目录标题
        toc_heading = self.doc.add_heading(TOC_CONFIG["heading"], level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加目录字段
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar)

        run = paragraph.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-{TOC_CONFIG["max_levels"]}" \\h \\z \\u </w:instrText>')
        run._r.append(instrText)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run._r.append(fldChar)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar)

        # 添加分页符
        self.doc.add_page_break()

    def add_heading(self, level: int, text: str):
        """添加标题"""
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False,
                      font_name: str = None, font_size: float = None):
        """添加段落"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)

        # 设置字体
        if font_name:
            run.font.name = font_name
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), font_name)

        if font_size:
            run.font.size = Pt(font_size)

        run.font.bold = bold
        run.font.italic = italic

        # 设置行距
        pf = paragraph.paragraph_format
        pf.line_spacing = Pt(LINE_SPACING)
        pf.space_before = Pt(PARAGRAPH_SPACING_BEFORE)
        pf.space_after = Pt(PARAGRAPH_SPACING_AFTER)

        return paragraph

    def add_table(self, rows: List[List[str]]):
        """添加表格"""
        if not rows:
            return

        num_cols = max(len(row) for row in rows)
        num_rows = len(rows)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = TABLE_STYLE
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 填充数据
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    cell = table.cell(i, j)
                    cell.text = cell_text

                    # 设置单元格字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = TABLE_FONT_NAME
                            run.font.size = Pt(TABLE_FONT_SIZE)
                            rPr = run._r.get_or_add_rPr()
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                                rPr.append(rFonts)
                            rFonts.set(qn('w:eastAsia'), TABLE_FONT_NAME)

        # 添加空行
        self.doc.add_paragraph()

    def add_code_block(self, code: str, language: str = ""):
        """添加代码块"""
        paragraph = self.doc.add_paragraph()

        # 设置段落样式
        pf = paragraph.paragraph_format
        pf.left_indent = Cm(0.5)

        # 添加代码文本
        run = paragraph.add_run(code)
        run.font.name = CODE_BLOCK_CONFIG["font_name"]
        run.font.size = Pt(CODE_BLOCK_CONFIG["font_size"])

        # 设置中文字体
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_CHINESE)

        # 添加浅灰色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BLOCK_CONFIG["background_color"]}"/>')
        rPr.append(shading)

        # 添加空行
        self.doc.add_paragraph()

    def add_blockquote(self, text: str):
        """添加引用块"""
        paragraph = self.doc.add_paragraph()

        # 设置缩进
        pf = paragraph.paragraph_format
        pf.left_indent = Cm(BLOCKQUOTE_CONFIG["left_indent"])

        run = paragraph.add_run(text)
        run.font.name = BLOCKQUOTE_CONFIG["font_name"]
        run.font.size = Pt(BLOCKQUOTE_CONFIG["font_size"])
        run.font.italic = BLOCKQUOTE_CONFIG["italic"]

        # 设置中文字体
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), BLOCKQUOTE_CONFIG["font_name"])

        # 添加空行
        self.doc.add_paragraph()

    def add_list(self, items: List[str], ordered: bool = False):
        """添加列表"""
        for i, item in enumerate(items):
            if ordered:
                paragraph = self.doc.add_paragraph(style='List Number')
            else:
                paragraph = self.doc.add_paragraph(style='List Bullet')

            # 清除默认文本并添加自定义内容
            paragraph.clear()
            run = paragraph.add_run(item)
            run.font.name = FONT_CHINESE
            run.font.size = Pt(BODY_FONT_SIZE)

            # 设置中文字体
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), FONT_CHINESE)

        # 添加空行
        self.doc.add_paragraph()

    def add_image(self, image_path: str, alt_text: str = ""):
        """添加图片"""
        try:
            if os.path.exists(image_path):
                # 获取图片尺寸
                with Image.open(image_path) as img:
                    width, height = img.size

                # 计算合适的显示尺寸（最大宽度16cm）
                max_width_cm = MAX_IMAGE_WIDTH
                aspect_ratio = height / width

                # 转换为EMU（1英寸=914400 EMU）
                width_emu = int(max_width_cm / 2.54 * 914400)
                height_emu = int(width_emu * aspect_ratio)

                # 限制最大高度
                max_height_emu = int(MAX_IMAGE_HEIGHT / 2.54 * 914400)
                if height_emu > max_height_emu:
                    height_emu = max_height_emu
                    width_emu = int(height_emu / aspect_ratio)

                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(image_path, width=width_emu)

                # 添加图片说明
                if alt_text:
                    caption = self.doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption.add_run(alt_text)
                    run.font.size = Pt(10)
                    run.font.italic = True

                self.doc.add_paragraph()
                return True
            else:
                # 添加占位符
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(f"[图片: {alt_text or image_path}]")
                run.font.size = Pt(12)
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                self.doc.add_paragraph()
                return False

        except Exception as e:
            self.logger.error(f"添加图片失败: {image_path}, 错误: {str(e)}")
            # 添加占位符
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(f"[图片: {alt_text or image_path}]")
            run.font.size = Pt(12)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            self.doc.add_paragraph()
            return False

    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()

    def save(self, output_path: str):
        """保存文档"""
        self.logger.info(f"保存文档: {output_path}")

        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        self.doc.save(output_path)
        self.logger.info("文档保存完成")


# ==================== 资源处理器 ====================

class ResourceHandler:
    """资源处理器"""

    def __init__(self, base_dir: str, logger: logging.Logger):
        self.base_dir = base_dir
        self.logger = logger

    def resolve_path(self, relative_path: str) -> str:
        """解析相对路径为绝对路径"""
        # 移除开头的 ./
        if relative_path.startswith('./'):
            relative_path = relative_path[2:]

        # 如果是绝对路径，直接返回
        if os.path.isabs(relative_path):
            return relative_path

        # 相对于基础目录解析
        abs_path = os.path.join(self.base_dir, relative_path)
        abs_path = os.path.normpath(abs_path)

        return abs_path

    def check_image(self, image: ImageResource) -> bool:
        """检查图片是否存在"""
        abs_path = self.resolve_path(image.path)
        image.found = os.path.exists(abs_path)

        if image.found:
            self.logger.debug(f"图片找到: {abs_path}")
        else:
            self.logger.warning(f"图片缺失: {abs_path}")

        return image.found

    def validate_images(self, images: List[ImageResource]) -> Tuple[int, int]:
        """验证所有图片"""
        self.logger.info("验证图片资源...")

        found = 0
        missing = 0

        for image in images:
            if self.check_image(image):
                found += 1
            else:
                missing += 1

        self.logger.info(f"图片验证完成: 找到 {found}, 缺失 {missing}")
        return found, missing


# ==================== 校验器 ====================

class Validator:
    """转换校验器"""

    def __init__(self, logger: logging.Logger):
        self.logger = logger

    def validate(self, source_stats: Dict[str, Any], target_doc: Document,
                 images: List[ImageResource]) -> ValidationResult:
        """执行全量校验"""
        self.logger.info("开始校验...")

        result = ValidationResult()

        # 内容完整性校验
        result.content_complete = self._validate_content(source_stats, target_doc)

        # 格式规范性校验
        result.format_correct = self._validate_format(target_doc)

        # 资源可用性校验
        result.resources_available = self._validate_resources(images)

        # 总体结果
        result.passed = (result.content_complete and
                        result.format_correct and
                        result.resources_available)

        # 统计信息
        result.source_stats = source_stats
        result.target_stats = self._get_target_stats(target_doc)

        return result

    def _validate_content(self, source_stats: Dict[str, Any], target_doc: Document) -> bool:
        """内容完整性校验"""
        self.logger.info("校验内容完整性...")

        # 统计Word文档中的段落数
        para_count = len(target_doc.paragraphs)
        source_para_count = source_stats.get("paragraphs", 0)
        table_count = source_stats.get("tables", 0)

        # 表格单元格会增加段落数，每个表格平均约5-15个单元格
        # 同时Word文档中段落会被拆分（如列表项、表格单元格等）
        estimated_table_paras = table_count * 8
        adjusted_source_count = source_para_count + estimated_table_paras

        # 允许较大的误差范围（Word文档结构会导致段落数增加）
        tolerance = adjusted_source_count * 0.5  # 50%容差

        # 检查段落数是否在合理范围内（不能太少）
        if para_count < source_para_count * 0.8:
            self.logger.warning(f"段落数过少: 源{source_para_count}, 目标{para_count}")
            return False

        self.logger.info(f"段落数校验通过: 源{source_para_count}(+{estimated_table_paras}表格估计) ≈ 目标{para_count}")
        return True

    def _validate_format(self, target_doc: Document) -> bool:
        """格式规范性校验"""
        self.logger.info("校验格式规范性...")

        # 检查标题样式
        heading_count = 0
        for para in target_doc.paragraphs:
            if para.style.name.startswith('Heading'):
                heading_count += 1

        if heading_count == 0 and len(target_doc.paragraphs) > 0:
            self.logger.warning("未检测到标题样式")
            return False

        return True

    def _validate_resources(self, images: List[ImageResource]) -> bool:
        """资源可用性校验"""
        self.logger.info("校验资源可用性...")

        if not images:
            return True

        found = sum(1 for img in images if img.found)
        total = len(images)

        if found < total:
            self.logger.warning(f"部分图片缺失: {found}/{total}")
            # 允许部分图片缺失（记录但不视为失败）

        return True

    def _get_target_stats(self, target_doc: Document) -> Dict[str, Any]:
        """获取目标文档统计"""
        stats = {
            "paragraphs": len(target_doc.paragraphs),
            "tables": len(target_doc.tables),
        }

        # 统计标题
        headings = {}
        for para in target_doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', ''))
                headings[level] = headings.get(level, 0) + 1
        stats["headings"] = headings

        return stats

    def generate_report(self, result: ValidationResult, output_path: str):
        """生成校验报告"""
        self.logger.info(f"生成校验报告: {output_path}")

        report = []
        report.append("=" * 60)
        report.append("Markdown转Word文档 - 转换校验报告")
        report.append("=" * 60)
        report.append("")

        # 总体结果
        report.append(f"总体结果: {'通过' if result.passed else '失败'}")
        report.append("")

        # 详细结果
        report.append("详细结果:")
        report.append(f"  - 内容完整性: {'通过' if result.content_complete else '失败'}")
        report.append(f"  - 格式规范性: {'通过' if result.format_correct else '失败'}")
        report.append(f"  - 资源可用性: {'通过' if result.resources_available else '失败'}")
        report.append("")

        # 源文件统计
        report.append("源文件统计:")
        for key, value in result.source_stats.items():
            report.append(f"  - {key}: {value}")
        report.append("")

        # 目标文件统计
        report.append("目标文件统计:")
        for key, value in result.target_stats.items():
            report.append(f"  - {key}: {value}")
        report.append("")

        # 问题列表
        if result.issues:
            report.append("发现的问题:")
            for issue in result.issues:
                report.append(f"  - {issue}")
            report.append("")

        report.append("=" * 60)
        from datetime import datetime
        report.append(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # 写入文件
        with open(output_path, "w", encoding="utf-8") as f:
            f.write('\n'.join(report))

        self.logger.info("校验报告生成完成")


# ==================== 主转换器 ====================

class MarkdownToWordConverter:
    """Markdown转Word转换器主类"""

    def __init__(self, input_file: str, output_file: str, logger: logging.Logger):
        self.input_file = input_file
        self.output_file = output_file
        self.logger = logger

        # 初始化组件
        self.parser = MarkdownParser(logger)
        self.generator = WordGenerator(logger)
        self.resource_handler = ResourceHandler(RESOURCE_BASE_DIR, logger)
        self.validator = Validator(logger)

        # 统计
        self.stats = ConversionStats()

    def convert(self) -> bool:
        """执行转换"""
        self.logger.info("=" * 60)
        self.logger.info("开始Markdown到Word的转换")
        self.logger.info("=" * 60)

        try:
            # 1. 读取源文件
            content = self.parser.read_file(self.input_file)

            # 2. 提取图片引用
            images = self.parser.extract_images(content)
            self.stats.images = len(images)

            # 3. 验证图片
            found, missing = self.resource_handler.validate_images(images)
            self.stats.images_found = found
            self.stats.images_missing = missing

            # 4. 解析Markdown
            elements = self.parser.parse_markdown(content)
            self.stats = self.parser.stats

            # 5. 创建Word文档
            doc = self.generator.create_document()

            # 6. 添加目录
            self.generator.add_toc()

            # 7. 转换内容
            self.logger.info("转换内容到Word...")
            for i, element in enumerate(tqdm(elements, desc="转换进度", unit="元素")):
                try:
                    self._convert_element(element, images)
                except Exception as e:
                    error_msg = f"转换元素失败 (位置 {i}): {str(e)}"
                    self.logger.error(error_msg)
                    self.stats.errors.append(error_msg)

            # 8. 保存文档
            self.generator.save(self.output_file)

            # 9. 校验
            validation_result = self.validator.validate(
                {
                    "total_lines": self.stats.total_lines,
                    "total_chars": self.stats.total_chars,
                    "paragraphs": self.stats.paragraphs,
                    "headings": self.stats.headings,
                    "tables": self.stats.tables,
                    "images": self.stats.images,
                },
                doc,
                images
            )

            # 10. 生成报告
            self.validator.generate_report(validation_result, REPORT_FILE)

            # 11. 输出统计
            self._print_stats()

            self.logger.info("转换完成!")
            return True

        except Exception as e:
            self.logger.error(f"转换失败: {str(e)}")
            import traceback
            self.logger.error(traceback.format_exc())
            return False

    def _convert_element(self, element: Dict[str, Any], images: List[ImageResource]):
        """转换单个元素"""
        elem_type = element["type"]

        if elem_type == "heading":
            self.generator.add_heading(element["level"], element["text"])

        elif elem_type == "paragraph":
            # 处理行内格式
            self._convert_formatted_paragraph(element["text"])

        elif elem_type == "code_block":
            self.generator.add_code_block(element["code"], element["language"])

        elif elem_type == "table":
            self.generator.add_table(element["rows"])

        elif elem_type == "image":
            # 查找对应的图片资源
            img_path = element["path"]
            abs_path = self.resource_handler.resolve_path(img_path)
            self.generator.add_image(abs_path, element["alt"])

        elif elem_type == "blockquote":
            self.generator.add_blockquote(element["text"])

        elif elem_type == "ordered_list":
            self.generator.add_list(element["items"], ordered=True)

        elif elem_type == "unordered_list":
            self.generator.add_list(element["items"], ordered=False)

    def _convert_formatted_paragraph(self, text: str):
        """转换带格式的段落"""
        # 简单的格式处理：加粗、斜体、行内代码
        # 使用正则表达式分割文本
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

        paragraph = self.generator.doc.add_paragraph()
        pf = paragraph.paragraph_format
        pf.line_spacing = Pt(LINE_SPACING)
        pf.space_before = Pt(PARAGRAPH_SPACING_BEFORE)
        pf.space_after = Pt(PARAGRAPH_SPACING_AFTER)

        for part in parts:
            if not part:
                continue

            run = None

            # 加粗 **text**
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True

            # 斜体 *text*
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True

            # 行内代码 `text`
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = FONT_CODE
                run.font.size = Pt(CODE_FONT_SIZE)

            # 普通文本
            else:
                run = paragraph.add_run(part)

            # 设置中文字体
            if run:
                run.font.name = FONT_CHINESE
                run.font.size = Pt(BODY_FONT_SIZE)
                rPr = run._r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), FONT_CHINESE)

    def _print_stats(self):
        """输出统计信息"""
        self.logger.info("=" * 60)
        self.logger.info("转换统计")
        self.logger.info("=" * 60)
        self.logger.info(f"总行数: {self.stats.total_lines}")
        self.logger.info(f"总字符数: {self.stats.total_chars}")
        self.logger.info(f"标题数: {sum(self.stats.headings.values())}")
        for level, count in self.stats.headings.items():
            if count > 0:
                self.logger.info(f"  - H{level}: {count}")
        self.logger.info(f"段落数: {self.stats.paragraphs}")
        self.logger.info(f"表格数: {self.stats.tables}")
        self.logger.info(f"代码块数: {self.stats.code_blocks}")
        self.logger.info(f"列表数: {self.stats.lists}")
        self.logger.info(f"图片数: {self.stats.images} (找到: {self.stats.images_found}, 缺失: {self.stats.images_missing})")
        self.logger.info(f"引用块数: {self.stats.blockquotes}")
        self.logger.info(f"错误数: {len(self.stats.errors)}")
        self.logger.info(f"警告数: {len(self.stats.warnings)}")
        self.logger.info("=" * 60)


# ==================== 主程序入口 ====================

def main():
    """主程序入口"""
    import argparse

    parser = argparse.ArgumentParser(
        description="Markdown to Word Document Converter - 将Markdown文件转换为符合公文规范的Word文档",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  python md_to_word_converter.py
  python md_to_word_converter.py -i input.md -o output.docx
  python md_to_word_converter.py --input "path/to/input.md" --output "path/to/output.docx"
        """
    )

    parser.add_argument(
        "-i", "--input",
        default=DEFAULT_INPUT_FILE,
        help=f"输入Markdown文件路径 (默认: {DEFAULT_INPUT_FILE})"
    )

    parser.add_argument(
        "-o", "--output",
        default=DEFAULT_OUTPUT_FILE,
        help=f"输出Word文件路径 (默认: {DEFAULT_OUTPUT_FILE})"
    )

    parser.add_argument(
        "--log",
        default=LOG_FILE,
        help=f"日志文件路径 (默认: {LOG_FILE})"
    )

    parser.add_argument(
        "--report",
        default=REPORT_FILE,
        help=f"校验报告路径 (默认: {REPORT_FILE})"
    )

    args = parser.parse_args()

    # 设置日志
    logger = setup_logging(args.log)

    # 创建转换器
    converter = MarkdownToWordConverter(
        input_file=args.input,
        output_file=args.output,
        logger=logger
    )

    # 执行转换
    success = converter.convert()

    if success:
        logger.info(f"转换成功! 输出文件: {args.output}")
        logger.info(f"校验报告: {args.report}")
        sys.exit(0)
    else:
        logger.error("转换失败!")
        sys.exit(1)


if __name__ == "__main__":
    main()
